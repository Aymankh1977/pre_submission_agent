import streamlit as st
import os
import base64
import json
import hashlib
import datetime
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv
from anthropic import Anthropic, NotFoundError
from pypdf import PdfReader

try:
    from duckduckgo_search import DDGS
    DDG_AVAILABLE = True
except ImportError:
    DDG_AVAILABLE = False

# ─── CONFIG ───────────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="DentEdTech™ | HPE Expert Reviewer",
    page_icon="🎓",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ─── CUSTOM STYLING ───────────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Playfair+Display:ital,wght@0,400;0,600;1,400&family=Source+Serif+4:ital,opsz,wght@0,8..60,300;0,8..60,400;1,8..60,300&family=JetBrains+Mono:wght@400;500&display=swap');

/* ── Root palette ── */
:root {
    --navy:    #0f2535;
    --navy2:   #1a3a4a;
    --teal:    #1d6b52;
    --teal2:   #2d8f6f;
    --gold:    #b5903a;
    --ivory:   #f8f5ef;
    --ivory2:  #f0ebe0;
    --stone:   #e8e2d6;
    --text:    #1a1612;
    --text2:   #4a4540;
    --text3:   #8a847a;
    --white:   #ffffff;
    --success: #1d6b52;
    --warn:    #8a5a00;
    --danger:  #7a2020;
}

/* ── Base app background ── */
.stApp {
    background: var(--ivory) !important;
    font-family: 'Source Serif 4', Georgia, serif !important;
}

/* ── Hide default Streamlit chrome — keep sidebar toggle visible ── */
#MainMenu, footer { visibility: hidden !important; }
.stDeployButton { display: none !important; }
/* Keep the header visible but make it transparent so toggle button works */
header[data-testid="stHeader"] {
    background: transparent !important;
    height: 0 !important;
    min-height: 0 !important;
}
/* Style the sidebar collapse button to match our navy theme */
[data-testid="collapsedControl"],
button[kind="header"] {
    color: #0f2535 !important;
    background: var(--ivory) !important;
}

/* ── Main content area ── */
.main .block-container {
    padding: 2rem 2.5rem 3rem !important;
    max-width: 1100px !important;
}

/* ── Sidebar ── */
[data-testid="stSidebar"] {
    background: #2c4a5a !important;
    border-right: 1px solid rgba(255,255,255,0.08) !important;
}

/* Section label text */
[data-testid="stSidebar"] p,
[data-testid="stSidebar"] span:not([data-baseweb]),
[data-testid="stSidebar"] .stMarkdown p {
    color: #c8d8e0 !important;
    font-family: 'Source Serif 4', Georgia, serif !important;
    font-size: 0.84rem !important;
    line-height: 1.6 !important;
}

/* Headings */
[data-testid="stSidebar"] h1 {
    font-family: 'Playfair Display', Georgia, serif !important;
    font-size: 1.35rem !important;
    font-weight: 600 !important;
    color: #eef4f8 !important;
    letter-spacing: 0.01em !important;
    margin-bottom: 0 !important;
}
[data-testid="stSidebar"] h2,
[data-testid="stSidebar"] h3 {
    font-family: 'Source Serif 4', serif !important;
    font-size: 0.72rem !important;
    font-weight: 400 !important;
    color: #7a9aaa !important;
    letter-spacing: 0.14em !important;
    text-transform: uppercase !important;
    margin: 0.8rem 0 0.4rem !important;
}

/* Captions */
[data-testid="stSidebar"] .stCaption p {
    color: #6a8898 !important;
    font-size: 0.73rem !important;
    font-style: italic !important;
    font-family: 'Source Serif 4', serif !important;
}

/* Dividers */
[data-testid="stSidebar"] hr {
    border: none !important;
    border-top: 1px solid rgba(255,255,255,0.1) !important;
    margin: 0.75rem 0 !important;
}

/* Checkbox labels */
[data-testid="stSidebar"] .stCheckbox label p,
[data-testid="stSidebar"] .stCheckbox label {
    color: #b8ccd8 !important;
    font-family: 'Source Serif 4', serif !important;
    font-size: 0.79rem !important;
    font-weight: 300 !important;
    line-height: 1.5 !important;
}

/* Selectbox label */
[data-testid="stSidebar"] .stSelectbox label p {
    color: #8aacbe !important;
    font-size: 0.74rem !important;
    letter-spacing: 0.08em !important;
    text-transform: uppercase !important;
    font-family: 'Source Serif 4', serif !important;
}

/* Selectbox control */
[data-testid="stSidebar"] [data-baseweb="select"] > div {
    background: rgba(255,255,255,0.08) !important;
    border: 1px solid rgba(255,255,255,0.15) !important;
    border-radius: 3px !important;
}
[data-testid="stSidebar"] [data-baseweb="select"] span,
[data-testid="stSidebar"] [data-baseweb="select"] div {
    color: #ddeaf0 !important;
    font-family: 'Source Serif 4', serif !important;
    font-size: 0.83rem !important;
}

/* File uploader */
[data-testid="stSidebar"] [data-testid="stFileUploader"] {
    background: rgba(255,255,255,0.05) !important;
    border: 1px dashed rgba(255,255,255,0.2) !important;
    border-radius: 3px !important;
}
[data-testid="stSidebar"] [data-testid="stFileUploader"] p,
[data-testid="stSidebar"] [data-testid="stFileUploader"] span {
    color: #9ab8c8 !important;
    font-size: 0.8rem !important;
}
[data-testid="stSidebar"] [data-testid="stFileUploader"] button {
    color: #b8d4e0 !important;
    border-color: rgba(255,255,255,0.2) !important;
    background: rgba(255,255,255,0.07) !important;
}

/* Upload label */
[data-testid="stSidebar"] .stFileUploader label p {
    color: #8aacbe !important;
    font-size: 0.74rem !important;
    letter-spacing: 0.08em !important;
    text-transform: uppercase !important;
}

/* Sidebar primary button — Run Analysis */
[data-testid="stSidebar"] .stButton:nth-child(1) > button,
[data-testid="stSidebar"] .stButton > button[kind="primary"] {
    background: rgba(29, 107, 82, 0.35) !important;
    border: 1px solid rgba(45, 143, 111, 0.5) !important;
    color: #90dfc0 !important;
    font-family: 'Source Serif 4', serif !important;
    font-size: 0.84rem !important;
    letter-spacing: 0.05em !important;
    border-radius: 3px !important;
    padding: 0.5rem 1rem !important;
    transition: all 0.2s ease !important;
    width: 100% !important;
}
[data-testid="stSidebar"] .stButton > button[kind="primary"]:hover {
    background: rgba(29, 107, 82, 0.55) !important;
    color: #c0f0dc !important;
}

/* All other sidebar buttons */
[data-testid="stSidebar"] .stButton > button {
    background: rgba(255,255,255,0.07) !important;
    border: 1px solid rgba(255,255,255,0.14) !important;
    color: #c8dde8 !important;
    font-family: 'Source Serif 4', serif !important;
    font-size: 0.82rem !important;
    letter-spacing: 0.04em !important;
    border-radius: 3px !important;
    padding: 0.45rem 1rem !important;
    transition: all 0.2s ease !important;
}
[data-testid="stSidebar"] .stButton > button:hover {
    background: rgba(255,255,255,0.13) !important;
    border-color: rgba(255,255,255,0.25) !important;
    color: #eef4f8 !important;
}

/* Progress bar */
[data-testid="stSidebar"] .stProgress > div {
    background: rgba(255,255,255,0.1) !important;
    border-radius: 2px !important;
}
[data-testid="stSidebar"] .stProgress > div > div {
    background: #3d9e7a !important;
    border-radius: 2px !important;
}
[data-testid="stSidebar"] .stProgress p {
    color: #7aacbe !important;
    font-size: 0.72rem !important;
}

/* Expander in sidebar */
[data-testid="stSidebar"] .streamlit-expanderHeader {
    background: rgba(255,255,255,0.06) !important;
    border: 1px solid rgba(255,255,255,0.1) !important;
    border-radius: 3px !important;
    color: #a8c4d0 !important;
    font-family: 'Source Serif 4', serif !important;
    font-size: 0.8rem !important;
    letter-spacing: 0.04em !important;
    padding: 0.5rem 0.8rem !important;
}
[data-testid="stSidebar"] .streamlit-expanderHeader:hover {
    background: rgba(255,255,255,0.1) !important;
    color: #d8eaf4 !important;
}
[data-testid="stSidebar"] .streamlit-expanderContent {
    background: rgba(255,255,255,0.03) !important;
    border: 1px solid rgba(255,255,255,0.07) !important;
    border-top: none !important;
    border-radius: 0 0 3px 3px !important;
}

/* Sidebar metrics */
[data-testid="stSidebar"] [data-testid="stMetric"] {
    background: rgba(255,255,255,0.06) !important;
    border: 1px solid rgba(255,255,255,0.1) !important;
    border-radius: 3px !important;
}
[data-testid="stSidebar"] [data-testid="stMetricLabel"] p {
    color: #6a8898 !important;
    font-size: 0.68rem !important;
    letter-spacing: 0.1em !important;
}
[data-testid="stSidebar"] [data-testid="stMetricValue"] {
    color: #cce0ea !important;
    font-family: 'Playfair Display', serif !important;
    font-size: 1.3rem !important;
}

/* Sidebar markdown tables */
[data-testid="stSidebar"] table {
    font-size: 0.75rem !important;
    color: #a8c0cc !important;
    border-color: rgba(255,255,255,0.1) !important;
}
[data-testid="stSidebar"] td, [data-testid="stSidebar"] th {
    color: #a8c0cc !important;
    border-color: rgba(255,255,255,0.08) !important;
    padding: 0.25rem 0.5rem !important;
}

/* Sidebar alerts */
[data-testid="stSidebar"] .stAlert {
    background: rgba(255,255,255,0.06) !important;
    border-radius: 3px !important;
}
[data-testid="stSidebar"] .stAlert p {
    color: #c8d8e0 !important;
    font-size: 0.78rem !important;
}

/* ── Main headings ── */
h1, h2, h3 {
    font-family: 'Playfair Display', Georgia, serif !important;
    color: var(--navy) !important;
    letter-spacing: -0.01em !important;
}
h1 { font-size: 2rem !important; font-weight: 600 !important; }
h2 { font-size: 1.35rem !important; font-weight: 600 !important; }
h3 { font-size: 1.1rem !important; font-weight: 400 !important; font-style: italic !important; }

p, li, label { font-family: 'Source Serif 4', Georgia, serif !important; }

/* ── Tabs ── */
.stTabs [data-baseweb="tab-list"] {
    background: transparent !important;
    border-bottom: 1px solid var(--stone) !important;
    gap: 0 !important;
}
.stTabs [data-baseweb="tab"] {
    font-family: 'Source Serif 4', serif !important;
    font-size: 0.82rem !important;
    font-weight: 400 !important;
    letter-spacing: 0.08em !important;
    text-transform: uppercase !important;
    color: var(--text3) !important;
    padding: 0.7rem 1.4rem !important;
    border-bottom: 2px solid transparent !important;
    background: transparent !important;
}
.stTabs [aria-selected="true"] {
    color: var(--navy) !important;
    border-bottom: 2px solid var(--gold) !important;
    font-weight: 400 !important;
}
.stTabs [data-baseweb="tab-panel"] {
    padding-top: 1.5rem !important;
}

/* ── Expanders ── */
.streamlit-expanderHeader {
    font-family: 'Source Serif 4', serif !important;
    font-size: 0.9rem !important;
    font-weight: 400 !important;
    color: var(--navy2) !important;
    background: var(--white) !important;
    border: 1px solid var(--stone) !important;
    border-radius: 4px !important;
    padding: 0.6rem 1rem !important;
    letter-spacing: 0.01em !important;
}
.streamlit-expanderHeader:hover {
    background: var(--ivory2) !important;
}
.streamlit-expanderContent {
    border: 1px solid var(--stone) !important;
    border-top: none !important;
    border-radius: 0 0 4px 4px !important;
    background: var(--white) !important;
    padding: 1rem 1.2rem !important;
}

/* ── Metrics ── */
[data-testid="stMetric"] {
    background: var(--white) !important;
    border: 1px solid var(--stone) !important;
    border-radius: 4px !important;
    padding: 0.75rem 1rem !important;
}
[data-testid="stMetricLabel"] {
    font-family: 'Source Serif 4', serif !important;
    font-size: 0.72rem !important;
    letter-spacing: 0.1em !important;
    text-transform: uppercase !important;
    color: var(--text3) !important;
}
[data-testid="stMetricValue"] {
    font-family: 'Playfair Display', serif !important;
    font-size: 1.6rem !important;
    color: var(--navy) !important;
}

/* ── Buttons (main content) ── */
.stButton > button {
    font-family: 'Source Serif 4', serif !important;
    font-size: 0.85rem !important;
    letter-spacing: 0.06em !important;
    border-radius: 3px !important;
    border: 1px solid var(--stone) !important;
    background: var(--white) !important;
    color: var(--navy2) !important;
    padding: 0.5rem 1.2rem !important;
    transition: all 0.18s ease !important;
}
.stButton > button:hover {
    background: var(--ivory2) !important;
    border-color: var(--navy2) !important;
}

/* ── Download buttons ── */
.stDownloadButton > button {
    font-family: 'Source Serif 4', serif !important;
    font-size: 0.82rem !important;
    letter-spacing: 0.05em !important;
    border-radius: 3px !important;
    background: var(--navy) !important;
    color: #e8e0d0 !important;
    border: none !important;
    padding: 0.5rem 1.2rem !important;
    transition: background 0.18s ease !important;
}
.stDownloadButton > button:hover {
    background: var(--navy2) !important;
}

/* ── Text inputs ── */
.stTextInput input, .stSelectbox select,
[data-baseweb="select"] {
    font-family: 'Source Serif 4', serif !important;
    font-size: 0.88rem !important;
    border-radius: 3px !important;
    border: 1px solid var(--stone) !important;
    background: var(--white) !important;
    color: var(--text) !important;
}
.stTextInput input:focus {
    border-color: var(--navy2) !important;
    box-shadow: 0 0 0 2px rgba(26,58,74,0.1) !important;
}

/* ── Chat input ── */
.stChatInput textarea {
    font-family: 'Source Serif 4', serif !important;
    font-size: 0.88rem !important;
    border-radius: 3px !important;
    border: 1px solid var(--stone) !important;
}

/* ── Chat messages ── */
[data-testid="stChatMessage"] {
    background: var(--white) !important;
    border: 1px solid var(--stone) !important;
    border-radius: 4px !important;
    padding: 0.75rem 1rem !important;
}

/* ── Info / Warning / Success / Error boxes ── */
.stAlert {
    border-radius: 3px !important;
    font-family: 'Source Serif 4', serif !important;
    font-size: 0.87rem !important;
    border-left-width: 3px !important;
}
[data-testid="stInfoMessage"] {
    background: #f0f4f8 !important;
    border-left-color: var(--navy2) !important;
    color: var(--navy) !important;
}
[data-testid="stWarningMessage"] {
    background: #fdf6e3 !important;
    border-left-color: var(--gold) !important;
    color: #5a4010 !important;
}
[data-testid="stSuccessMessage"] {
    background: #f0f8f4 !important;
    border-left-color: var(--teal) !important;
    color: #0f3a28 !important;
}
[data-testid="stErrorMessage"] {
    background: #fdf0f0 !important;
    border-left-color: var(--danger) !important;
}

/* ── Progress bar ── */
.stProgress > div > div {
    background: var(--teal) !important;
    border-radius: 2px !important;
}
.stProgress > div {
    background: var(--stone) !important;
    border-radius: 2px !important;
}

/* ── Captions ── */
.stCaption, .stCaption p {
    font-family: 'Source Serif 4', serif !important;
    font-style: italic !important;
    color: var(--text3) !important;
    font-size: 0.78rem !important;
}

/* ── Markdown text ── */
.stMarkdown p {
    font-family: 'Source Serif 4', Georgia, serif !important;
    font-size: 0.92rem !important;
    line-height: 1.75 !important;
    color: var(--text2) !important;
}
.stMarkdown blockquote {
    border-left: 3px solid var(--gold) !important;
    background: var(--ivory2) !important;
    padding: 0.8rem 1.2rem !important;
    border-radius: 0 3px 3px 0 !important;
    font-style: italic !important;
    color: var(--text) !important;
}
.stMarkdown code {
    font-family: 'JetBrains Mono', monospace !important;
    font-size: 0.78rem !important;
    background: var(--ivory2) !important;
    color: var(--navy2) !important;
    padding: 0.1em 0.4em !important;
    border-radius: 2px !important;
}

/* ── Dividers ── */
hr {
    border: none !important;
    border-top: 1px solid var(--stone) !important;
    margin: 1.2rem 0 !important;
}

/* ── Selectbox ── */
[data-baseweb="select"] > div {
    border-radius: 3px !important;
    border-color: var(--stone) !important;
    background: var(--white) !important;
    font-family: 'Source Serif 4', serif !important;
    font-size: 0.85rem !important;
}

/* ── File uploader ── */
[data-testid="stFileUploader"] {
    border: 1px dashed var(--stone) !important;
    border-radius: 4px !important;
    background: var(--white) !important;
    padding: 0.5rem !important;
}
[data-testid="stFileUploader"]:hover {
    border-color: var(--navy2) !important;
    background: var(--ivory2) !important;
}

/* ── Status widget ── */
[data-testid="stStatusWidget"] {
    font-family: 'Source Serif 4', serif !important;
    font-size: 0.85rem !important;
    border-radius: 3px !important;
    border: 1px solid var(--stone) !important;
}

/* ── Spinner ── */
.stSpinner > div {
    border-top-color: var(--navy) !important;
}

/* ── Scrollbar ── */
::-webkit-scrollbar { width: 5px; height: 5px; }
::-webkit-scrollbar-track { background: var(--ivory2); }
::-webkit-scrollbar-thumb { background: var(--stone); border-radius: 3px; }
::-webkit-scrollbar-thumb:hover { background: var(--text3); }

/* ── Checkbox ── */
.stCheckbox [data-testid="stCheckbox"] {
    accent-color: var(--teal) !important;
}
</style>
""", unsafe_allow_html=True)

# ─── API KEY ──────────────────────────────────────────────────────────────────
try:
    api_key = st.secrets.get("ANTHROPIC_API_KEY")
except Exception:
    api_key = None

if not api_key:
    load_dotenv()
    api_key = os.getenv("ANTHROPIC_API_KEY")

if not api_key:
    st.error("🚨 ANTHROPIC_API_KEY is missing. Add it to `.env` or Streamlit Secrets.")
    st.stop()

client = Anthropic(api_key=api_key)

# ─── MODELS ───────────────────────────────────────────────────────────────────
PRIMARY_MODEL  = "claude-opus-4-5"
FALLBACK_MODEL = "claude-sonnet-4-5"
CHAT_MODEL     = "claude-haiku-4-5-20251001"

# Access gate removed — platform is open access
SESSION_ANALYSIS_CAP = 999

# ─── JOURNALS ─────────────────────────────────────────────────────────────────
JOURNALS = [
    # ISI / Web of Science Indexed
    "Medical Education (ISI)",
    "Academic Medicine (ISI)",
    "Medical Teacher (ISI)",
    "Advances in Health Sciences Education (ISI)",
    "Journal of Dental Research (ISI)",
    "Community Dentistry and Oral Epidemiology (ISI)",
    "European Journal of Oral Sciences (ISI)",
    "Journal of Dental Sciences (ISI)",
    # Scopus Indexed
    "BMC Medical Education (Scopus)",
    "JGME – Journal of Graduate Medical Education (Scopus)",
    "Teaching and Learning in Medicine (Scopus)",
    "International Journal of Medical Education – IJME (Scopus)",
    "GMS Journal for Medical Education (Scopus)",
    "Education for Health (Scopus)",
    "Journal of Educational Evaluation for Health Professions (Scopus)",
    "MedEdPublish (Scopus)",
    "Dentistry Journal – MDPI (Scopus)",
    # Dental Education Specialist
    "Journal of Dental Education (JDE)",
    "European Journal of Dental Education (EJDE)",
    "British Dental Journal (BDJ)",
    "Journal of Dentistry",
    "Dental Education Today (ADEE)",
    "Journal of Dental Hygiene Education",
]

# ─── STUDY DESIGNS ────────────────────────────────────────────────────────────
STUDY_DESIGNS = [
    "Auto-detect from manuscript",
    "Randomised Controlled Trial (RCT)",
    "Non-randomised / Quasi-experimental",
    "Cross-sectional Survey",
    "Systematic Review / Meta-analysis",
    "Qualitative Study",
    "Mixed Methods",
    "Case Study / Case Report",
    "Theoretical / Conceptual Paper",
    "Delphi Study",
    "Program Evaluation",
]

# Criteria that only apply to specific design families
QUANTITATIVE_ONLY   = {"sample_power", "instrument_validity"}
QUALITATIVE_ONLY    = {"positionality", "transferability"}
NOT_FOR_THEORETICAL = {"sample_power", "instrument_validity",
                       "positionality", "transferability", "consort_srqr"}

# ─── REVIEW CRITERIA ──────────────────────────────────────────────────────────
REVIEW_CRITERIA = {
    "research_question": "Research question clarity & PICO/SPIDER framing",
    "methodology":       "Methodology rigor & reproducibility",
    "consort_srqr":      "CONSORT / SRQR / COREQ guideline adherence",
    "kirkpatrick":       "Kirkpatrick level outcomes achieved",
    "citations":         "Citation currency, completeness & in-text accuracy",
    "statistics":        "Statistical / qualitative data analysis soundness",
    "ethics":            "Ethical considerations & positionality",
    "golden_thread":     "Golden thread coherence (RQ → method → results → conclusion)",
}

# ─── SESSION STATE ─────────────────────────────────────────────────────────────
# Increment this version string any time you want to force all sessions to reset
# (e.g. after a major deployment that adds new gates or changes access logic)
APP_VERSION = "3.0.0"

defaults = {
    "access_granted":        True,
    "access_partner":        "Open Access",
    "is_admin":              True,
    "consent_given":         False,
    "analyses_this_session": 0,
    "pdf_base64":            None,
    "pdf_name":              "",
    "pdf_hash":              "",
    "pdf_text":              "",
    "report":                None,
    "raw_report":            "",
    "chat_history":          [],
    "model_used":            "",
    "session_start":         None,
    "upload_count":          0,
    "similarity_report":     None,
    "raw_similarity":        "",
    "search_results":        [],
    "agreement_log":         [],
    "feedback_given":        False,
    "feedback_submitted":    False,
    "_app_version":          APP_VERSION,
}

# Force full session reset if app version has changed
if st.session_state.get("_app_version") != APP_VERSION:
    for k in list(st.session_state.keys()):
        del st.session_state[k]

for k, v in defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v

if st.session_state.session_start is None:
    st.session_state.session_start = datetime.datetime.utcnow()

# Gate 1 (access code) removed — platform is open access

# ─── GATE 2: CONSENT ──────────────────────────────────────────────────────────
if not st.session_state.consent_given:
    st.markdown(
        f"""
        <div style="max-width:600px;margin:5rem auto 2rem;">
          <div style="font-family:'Playfair Display',Georgia,serif;
                      font-size:0.7rem;letter-spacing:0.3em;text-transform:uppercase;
                      color:#8a9aaa;margin-bottom:1rem;text-align:center;">
            Health Professions Education
          </div>
          <h1 style="font-family:'Playfair Display',Georgia,serif;font-size:2.2rem;
                     font-weight:600;color:#0f2535;margin:0 0 0.2rem;
                     letter-spacing:-0.02em;text-align:center;">
            DentEdTech™
          </h1>
          <div style="font-family:'Source Serif 4',Georgia,serif;font-size:0.9rem;
                      font-style:italic;color:#6a7a8a;text-align:center;margin-bottom:0.8rem;">
            Manuscript Intelligence Platform
          </div>
          <div style="width:40px;height:1px;background:#b5903a;margin:0 auto 2rem;"></div>
          <div style="background:#f8f5ef;border:1px solid #e8e2d6;border-radius:4px;
                      padding:1.5rem 1.8rem;margin-bottom:1.5rem;">
            <div style="font-family:'Playfair Display',serif;font-size:1rem;
                        font-weight:600;color:#0f2535;margin-bottom:1rem;">
              Data &amp; Confidentiality Notice
            </div>
            <ul style="font-family:'Source Serif 4',Georgia,serif;font-size:0.86rem;
                       line-height:1.9;color:#4a4540;padding-left:1.2rem;margin:0;">
              <li>Manuscripts are transmitted to <strong>Anthropic's API</strong> for analysis.
                  Anthropic does <strong>not</strong> train models on API data.</li>
              <li>Documents are held <strong>in memory only</strong> — never written to disk.</li>
              <li>Your session clears automatically when you close the browser tab.</li>
              <li>Do <strong>not</strong> upload manuscripts containing patient-identifiable
                  data or material under confidentiality agreement.</li>
            </ul>
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        confirmed = st.checkbox(
            "I have read the data notice above and confirm the manuscript I will upload "
            "does not contain restricted or patient-identifiable data."
        )
        if st.button("Accept & Continue →", disabled=not confirmed, use_container_width=True):
            st.session_state.consent_given = True
            st.rerun()
    st.stop()

# ─── GATE 3: SESSION CAP ──────────────────────────────────────────────────────
def check_session_cap() -> bool:
    return True

def show_cap_warning():
    pass

# ─── HELPERS ──────────────────────────────────────────────────────────────────
def encode_pdf(uploaded_file) -> tuple[str, str, str]:
    raw = uploaded_file.read()
    b64 = base64.standard_b64encode(raw).decode("utf-8")
    sha = hashlib.sha256(raw).hexdigest()
    try:
        reader = PdfReader(BytesIO(raw))
        text = "\n".join(p.extract_text() or "" for p in reader.pages)
    except Exception:
        text = ""
    return b64, text, sha


def clear_session_data():
    sensitive_keys = [
        "pdf_base64", "pdf_name", "pdf_hash", "pdf_text",
        "report", "raw_report", "chat_history", "model_used",
        "similarity_report", "raw_similarity", "search_results",
        "feedback_given",
    ]
    for k in sensitive_keys:
        st.session_state[k] = defaults[k]
    st.session_state.upload_count = 0


def build_system_prompt(journal: str) -> str:
    import datetime as _dt
    today = _dt.datetime.now(_dt.timezone.utc).strftime("%B %Y")
    year  = _dt.datetime.now(_dt.timezone.utc).year
    return (
        f"You are a Senior Editor and double-blind Peer Reviewer for '{journal}', "
        "one of the most rigorous journals in Health Professions Education (HPE). "
        f"Today's date is {today}. "
        f"References more than 5 years old (before {year-5}) should be flagged unless seminal. "
        "Flag highly relevant papers from the last 2 years that are missing. "
        "Your reviews are precise, evidence-based, constructive, and specific. "
        "Quote exact passages to substantiate every criticism. Never fabricate content. "
        "REPORTING FRAMEWORKS: CONSORT for RCTs, SRQR for qualitative, "
        "COREQ for interviews/focus groups, PRISMA for systematic reviews. "
        "Assess outcomes through Kirkpatrick's four-level framework. "
        "STUDY DESIGN AWARENESS: First identify the study design. Apply only criteria appropriate "
        "to that design. Flag criteria as N/A where not applicable. "
        "Quantitative: require sample size justification and power calculation. "
        "Qualitative: require positionality, reflexivity, and trustworthiness criteria. "
        "Theoretical: assess conceptual rigour, not empirical criteria. "
        "WRITING QUALITY STANDARD (based on exemplar HPE papers): "
        "Accepted papers explicitly identify a knowledge gap (not just a topic), state the RQ formally, "
        "have reproducible methods, honest specific limitations, concrete practice implications, "
        "and an authentic scholarly voice — precise, specific, appropriately hedged, evidence-grounded. "
        "HUMANISATION: Flag generic/formulaic passages. "
        "Humanised academic writing uses active voice where appropriate, grounds claims in specific "
        "evidence, avoids hollow phrases ('it is worth noting', 'contributes to the literature'), "
        "and maintains a consistent scholarly voice. Benchmark: JDE, EJDE, Medical Teacher, Medical Education. "
        "GOLDEN THREAD: Scrutinise the logical chain from RQ through methods, results, conclusion. "
        "CONFIDENCE: When uncertain, flag explicitly with a specific explanation."
    )


def build_review_prompt(selected_criteria: list[str], journal: str, study_design: str = "Auto-detect from manuscript") -> str:
    # Build criteria block — mark design-specific criteria appropriately
    quant_designs    = {"Randomised Controlled Trial (RCT)", "Non-randomised / Quasi-experimental",
                        "Cross-sectional Survey", "Delphi Study", "Program Evaluation"}
    qual_designs     = {"Qualitative Study", "Mixed Methods"}
    theory_designs   = {"Theoretical / Conceptual Paper"}
    is_quant         = study_design in quant_designs
    is_qual          = study_design in qual_designs or study_design == "Auto-detect from manuscript"
    is_theory        = study_design in theory_designs

    criteria_lines = []
    for i, c in enumerate(selected_criteria):
        label = REVIEW_CRITERIA[c]
        if c in QUANTITATIVE_ONLY and not is_quant and study_design != "Auto-detect from manuscript":
            label += " [APPLY ONLY IF QUANTITATIVE — flag N/A otherwise]"
        elif c in QUALITATIVE_ONLY and not is_qual and study_design != "Auto-detect from manuscript":
            label += " [APPLY ONLY IF QUALITATIVE — flag N/A otherwise]"
        elif c in NOT_FOR_THEORETICAL and is_theory:
            label += " [NOT APPLICABLE for theoretical papers — flag N/A]"
        criteria_lines.append(f"  {i+1}. {label}")
    criteria_block = "\n".join(criteria_lines)

    design_instruction = ""
    if study_design != "Auto-detect from manuscript":
        design_instruction = f"\nCONFIRMED STUDY DESIGN: {study_design}. Apply criteria accordingly.\n"
    else:
        design_instruction = "\nSTUDY DESIGN: Auto-detect from the manuscript before applying criteria.\n"

    return f"""Perform a comprehensive peer review of this manuscript submitted to '{journal}'.
{design_instruction}
SELECTED REVIEW CRITERIA:
{criteria_block}

CONFIDENCE INSTRUCTION:
After forming your verdict, assess your own confidence. If the manuscript is borderline,
if the methodology is ambiguous, or if the verdict could reasonably go either way,
set confidence to Moderate or Low and explain in confidence_note.
Only set High if the verdict is clear and unambiguous.

Return ONLY a valid JSON object — no markdown fences, no preamble:

{{
  "detected_study_design": "<detected or confirmed study design>",
  "verdict": "Accept | Minor Revisions | Major Revisions | Reject",
  "overall_score": <integer 1-100>,
  "confidence": "High | Moderate | Low",
  "confidence_note": "<what is uncertain and why the editor should verify>",
  "executive_summary": "<2-3 sentence overall assessment>",
  "scores": {{
    "novelty": <1-10>,
    "methodology": <1-10>,
    "clarity": <1-10>,
    "citations": <1-10>,
    "ethics": <1-10>,
    "writing_quality": <1-10>,
    "humanisation": <1-10>
  }},
  "strengths": ["<strength>"],
  "weaknesses": [
    {{
      "section": "Abstract|Introduction|Methods|Results|Discussion|Citations|Writing",
      "issue": "<specific issue quoting the manuscript>",
      "severity": "major|minor",
      "suggestion": "<concrete fix>"
    }}
  ],
  "section_comments": {{
    "abstract": "<comment>",
    "introduction": "<gap identified? citations current? RQ explicit? theoretical framework named?>",
    "methods": "<reproducibility, guideline adherence, sample size, power calculation if quant>",
    "results": "<clarity, alignment with RQ>",
    "discussion": "<overstating? Kirkpatrick level? golden thread? implications concrete?>"
  }},
  "design_specific_assessment": {{
    "research_gap_statement": "<explicit gap or just topic?>",
    "study_design_justification": "<design justified for RQ?>",
    "sample_power": "<power calculation adequate | N/A>",
    "instrument_validity": "<validated tools cited | N/A>",
    "theoretical_framework": "<named framework applied or absent>",
    "positionality": "<declared | N/A for quantitative>",
    "transferability": "<trustworthiness addressed | N/A>",
    "limitations_quality": "<specific/honest or generic?>",
    "implications_concreteness": "<concrete educator actions stated?>"
  }},
  "writing_quality_assessment": {{
    "academic_register": "<tone appropriate and consistent?>",
    "humanisation_score": "<authentic voice, specific examples?>",
    "generic_phrases_found": ["<hollow phrase quoted>"],
    "humanisation_suggestions": ["<rewrite suggestion>"],
    "voice_consistency": "<consistent scholarly voice?>"
  }},
  "golden_thread": "<RQ to methodology to results to conclusion coherence>",
  "kirkpatrick_level": {{
    "level": <1|2|3|4|"N/A">,
    "justification": "<why — or why N/A>"
  }},
  "citation_audit": {{
    "missing_key_references": ["<Author Year — why relevant>"],
    "potentially_outdated": ["<citation — reason>"],
    "mismatches": "<issues or None identified>"
  }},
  "actionable_recommendations": ["<specific action>"],
  "author_feedback_summary": "<2-3 constructive sentences to authors — encouraging, no verdict>",
  "editor_note": "<confidential note to editor only>"
}}"""


def build_similarity_prompt(search_results: list[dict]) -> str:
    search_block = ""
    if search_results:
        search_block = "\n\nSIMILAR PUBLISHED PAPERS FOUND ONLINE:\n"
        for i, r in enumerate(search_results, 1):
            search_block += (
                f"\n[{i}] Title: {r['title']}\n"
                f"    URL: {r['url']}\n"
                f"    Summary: {r['body'][:300]}\n"
            )
    return f"""You are an academic integrity and publication similarity specialist.
Analyse this manuscript for originality and similarity risks.

1. SIMILARITY TO PUBLISHED LITERATURE: Using the search results below, assess overlap
   with already-published papers. For each similar paper explain what overlaps and
   what the authors should do to differentiate.

2. INTERNAL ORIGINALITY AUDIT: Identify passages that are boilerplate, contain
   factual claims with no citation, show style changes suggesting imported text,
   or repeat between sections.

3. METHODS SECTION RISK: Assess whether the methods reads as original or copied.

4. OVERALL RISK with specific rewrite advice.
{search_block}

Return ONLY a valid JSON object — no markdown, no preamble:

{{
  "overall_risk_level": "Low | Moderate | High | Very High",
  "estimated_similarity_risk_percent": <integer 0-100>,
  "disclaimer": "AI-based risk assessment only. Not equivalent to Turnitin or iThenticate. Always use your institution's official similarity checker before submission.",
  "similar_publications": [
    {{
      "title": "<title>",
      "url": "<url>",
      "overlap_description": "<what overlaps>",
      "overlap_type": "topic | methodology | findings | framing | significant overlap",
      "risk_level": "Low | Moderate | High",
      "recommendation": "<what authors should do>"
    }}
  ],
  "boilerplate_sections": [
    {{
      "section": "<where>",
      "passage": "<quoted text>",
      "risk": "<why high risk>",
      "suggestion": "<how to rewrite>"
    }}
  ],
  "citation_free_claims": [
    {{
      "passage": "<quoted text>",
      "risk": "<what needs citation>",
      "suggestion": "<what to cite>"
    }}
  ],
  "internal_repetition": [
    {{
      "passage": "<quoted text>",
      "appears_in": ["<section 1>", "<section 2>"]
    }}
  ],
  "methods_risk": "<assessment of methods originality>",
  "priority_rewrites": ["<specific rewrite instruction>"],
  "submission_readiness": "<overall verdict and advice>"
}}"""


def call_api_with_pdf(system: str, user_prompt: str, model: str, max_tok: int = 4096) -> str:
    response = client.messages.create(
        model=model,
        max_tokens=max_tok,
        system=system,
        messages=[{
            "role": "user",
            "content": [
                {
                    "type": "document",
                    "source": {
                        "type": "base64",
                        "media_type": "application/pdf",
                        "data": st.session_state.pdf_base64,
                    },
                },
                {"type": "text", "text": user_prompt},
            ],
        }],
    )
    return response.content[0].text


def call_api_with_text(system: str, user_prompt: str, model: str, max_tok: int = 4096) -> str:
    text = st.session_state.pdf_text[:100_000]
    full_prompt = f"MANUSCRIPT TEXT:\n{text}\n\n{user_prompt}"
    response = client.messages.create(
        model=model,
        max_tokens=max_tok,
        system=system,
        messages=[{"role": "user", "content": full_prompt}],
    )
    return response.content[0].text


def extract_search_queries(pdf_text: str) -> list[str]:
    prompt = (
        "Read this manuscript excerpt and extract exactly 5 short search queries "
        "(4-8 words each) representing the most distinctive claims, methods, or findings. "
        "Return ONLY a JSON array of 5 strings.\n\n"
        f"MANUSCRIPT EXCERPT:\n{pdf_text[:8000]}"
    )
    try:
        r = client.messages.create(
            model=CHAT_MODEL, max_tokens=300,
            messages=[{"role": "user", "content": prompt}],
        )
        raw = r.content[0].text.strip()
        return json.loads(raw[raw.index("["):raw.rindex("]")+1])
    except Exception:
        lines = [l.strip() for l in pdf_text.split("\n") if len(l.strip()) > 40]
        return [lines[0][:80]] if lines else ["dental education quality assurance AI"]


def search_web(queries: list[str]) -> list[dict]:
    if not DDG_AVAILABLE:
        return []
    results, seen = [], set()
    try:
        with DDGS() as ddgs:
            for q in queries:
                try:
                    for r in ddgs.text(
                        f"{q} site:pubmed.ncbi.nlm.nih.gov OR site:researchgate.net "
                        f"OR site:tandfonline.com OR site:wiley.com OR site:springer.com "
                        f"OR site:sciencedirect.com",
                        max_results=3,
                    ):
                        url = r.get("href", "")
                        if url not in seen:
                            seen.add(url)
                            results.append({
                                "title": r.get("title", ""),
                                "url": url,
                                "body": r.get("body", ""),
                                "query": q,
                            })
                except Exception:
                    continue
    except Exception:
        pass
    return results[:12]


def parse_json(raw: str) -> dict | None:
    """Parse JSON robustly — handles fences, preamble, truncation, trailing text."""
    if not raw:
        return None
    cleaned = raw.strip()

    # Strip markdown code fences
    if "```" in cleaned:
        for part in cleaned.split("```"):
            part = part.strip()
            if part.startswith("json"):
                part = part[4:].strip()
            if part.startswith("{"):
                cleaned = part
                break

    # Find start of JSON object
    try:
        start = cleaned.index("{")
    except ValueError:
        return None

    # Brace matching to find outermost closing brace
    depth, end = 0, start
    for i, ch in enumerate(cleaned[start:], start):
        if ch == "{":
            depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0:
                end = i + 1
                break

    candidate = cleaned[start:end]

    # If JSON was truncated (depth never closed), repair it
    if depth != 0:
        open_count = candidate.count("{") - candidate.count("}")
        open_arr   = candidate.count("[") - candidate.count("]")
        candidate  = candidate.rstrip().rstrip(",").rstrip()
        candidate += "]" * max(0, open_arr) + "}" * max(0, open_count)

    try:
        return json.loads(candidate)
    except json.JSONDecodeError:
        # Last resort: rindex approach
        try:
            s = cleaned.index("{")
            e = cleaned.rindex("}") + 1
            return json.loads(cleaned[s:e])
        except Exception:
            return None


def create_author_feedback_docx(report: dict) -> bytes:
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)
    h = doc.add_heading("DentEdTech™ — Manuscript Development Feedback", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp = doc.sections[0].footer.paragraphs[0]
    fp.text = "DentEdTech™ Author Feedback — For manuscript improvement only. Not an editorial decision."
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def h2(t): doc.add_heading(t, level=2)

    note = doc.add_paragraph()
    note.add_run(
        "This report is designed to help you strengthen your manuscript before submission. "
        "All feedback is constructive and focused on enhancing the quality of your work."
    ).italic = True

    if report.get("author_feedback_summary"):
        h2("Overall Feedback")
        doc.add_paragraph(report["author_feedback_summary"])

    scores = report.get("scores", {})
    if scores:
        h2("Dimension Scores")
        for k, v in scores.items():
            doc.add_paragraph(f"{k.capitalize()}: {v}/10", style="List Bullet")

    kp = report.get("kirkpatrick_level", {})
    if kp:
        h2("Educational Outcomes (Kirkpatrick Level)")
        doc.add_paragraph(
            f"Your manuscript demonstrates Level {kp.get('level','?')} outcomes. "
            f"{kp.get('justification','')}"
        )

    h2("Logical Coherence (Golden Thread)")
    doc.add_paragraph(report.get("golden_thread", ""))

    if report.get("strengths"):
        h2("Strengths to Build On")
        for s in report["strengths"]:
            doc.add_paragraph(s, style="List Bullet")

    if report.get("weaknesses"):
        h2("Areas Requiring Attention")
        for w in report["weaknesses"]:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"[{w.get('severity','').upper()} — {w.get('section','')}] ").bold = True
            p.add_run(w.get("issue", ""))
            if w.get("suggestion"):
                doc.add_paragraph(f"Suggested action: {w['suggestion']}", style="List Bullet")

    sc = report.get("section_comments", {})
    if sc:
        h2("Section-by-Section Feedback")
        for sec, comment in sc.items():
            p = doc.add_paragraph()
            p.add_run(sec.capitalize() + ": ").bold = True
            p.add_run(comment)

    ca = report.get("citation_audit", {})
    if ca:
        h2("Citation & Reference Feedback")
        for ref in ca.get("missing_key_references", []):
            doc.add_paragraph(f"Consider citing: {ref}", style="List Bullet")
        for ref in ca.get("potentially_outdated", []):
            doc.add_paragraph(f"Review currency of: {ref}", style="List Bullet")
        doc.add_paragraph(f"Reference list consistency: {ca.get('mismatches','None identified')}")

    if report.get("actionable_recommendations"):
        h2("Priority Actions Before Resubmission")
        for i, rec in enumerate(report["actionable_recommendations"], 1):
            doc.add_paragraph(f"{i}. {rec}")

    h2("Next Steps")
    doc.add_paragraph(
        "We encourage you to address the points above before resubmission. "
        "The research you are conducting is valuable to the field — "
        "these revisions will give it the best chance of acceptance."
    )

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def create_docx(report: dict | None, raw: str) -> bytes:
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)
    h = doc.add_heading("DentEdTech™ — HPE Peer Review Report", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp = doc.sections[0].footer.paragraphs[0]
    fp.text = (
        "CONFIDENTIAL — Generated by DentEdTech™. "
        "Powered by Anthropic API. Not for redistribution."
    )
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if report is None:
        doc.add_paragraph(raw)
    else:
        def h2(t): doc.add_heading(t, level=2)
        verdict    = report.get("verdict", "—")
        score      = report.get("overall_score", "—")
        confidence = report.get("confidence", "—")
        p = doc.add_paragraph()
        run = p.add_run(f"Verdict: {verdict}  |  Score: {score}/100  |  Confidence: {confidence}")
        run.bold = True; run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x4A)
        conf_note = report.get("confidence_note", "")
        if conf_note:
            cn = doc.add_paragraph()
            cn.add_run("Confidence Note: ").bold = True
            cn.add_run(conf_note).italic = True
        doc.add_paragraph(report.get("executive_summary", ""))
        h2("Dimension Scores")
        for k, v in report.get("scores", {}).items():
            doc.add_paragraph(f"{k.capitalize()}: {v}/10", style="List Bullet")
        kp = report.get("kirkpatrick_level", {})
        if kp:
            h2("Kirkpatrick Level")
            doc.add_paragraph(f"Level {kp.get('level','?')}: {kp.get('justification','')}")
        h2("Golden Thread Analysis")
        doc.add_paragraph(report.get("golden_thread", ""))
        h2("Strengths")
        for s in report.get("strengths", []):
            doc.add_paragraph(s, style="List Bullet")
        h2("Weaknesses")
        for w in report.get("weaknesses", []):
            doc.add_paragraph(
                f"[{w.get('severity','').upper()} — {w.get('section','')}] "
                f"{w.get('issue','')}\n→ {w.get('suggestion','')}",
                style="List Bullet",
            )
        h2("Section-by-Section Comments")
        for sec, comment in report.get("section_comments", {}).items():
            p = doc.add_paragraph()
            p.add_run(sec.capitalize() + ": ").bold = True
            p.add_run(comment)
        h2("Citation Audit")
        ca = report.get("citation_audit", {})
        for ref in ca.get("missing_key_references", []):
            doc.add_paragraph(f"Missing: {ref}", style="List Bullet")
        for ref in ca.get("potentially_outdated", []):
            doc.add_paragraph(f"Outdated: {ref}", style="List Bullet")
        doc.add_paragraph(f"Mismatches: {ca.get('mismatches','None identified')}")
        h2("Actionable Recommendations")
        for i, rec in enumerate(report.get("actionable_recommendations", []), 1):
            doc.add_paragraph(f"{i}. {rec}")
        h2("Confidential Note to Editor")
        p = doc.add_paragraph(report.get("editor_note", ""))
        for run in p.runs: run.italic = True

    buf = BytesIO(); doc.save(buf); return buf.getvalue()


def render_badge(text: str, colour_key: str, palette: dict) -> str:
    bg, fg = palette.get(colour_key.lower().strip(), ("#e2e3e5", "#383d41"))
    return (
        f'<span style="background:{bg};color:{fg};padding:4px 14px;'
        f'border-radius:20px;font-weight:600;font-size:0.9rem;">{text}</span>'
    )

VERDICT_COLOURS = {
    "accept": ("#d4edda", "#155724"),
    "minor":  ("#fff3cd", "#856404"),
    "major":  ("#f8d7da", "#721c24"),
    "reject": ("#f8d7da", "#491217"),
}
CONFIDENCE_COLOURS = {
    "high":     ("#d4edda", "#155724"),
    "moderate": ("#fff3cd", "#856404"),
    "low":      ("#f8d7da", "#721c24"),
}
RISK_COLOURS = {
    "low":       ("#d4edda", "#155724"),
    "moderate":  ("#fff3cd", "#856404"),
    "high":      ("#f8d7da", "#721c24"),
    "very high": ("#f5c6cb", "#491217"),
}


def verdict_key(v: str) -> str:
    vl = v.lower()
    if "reject" in vl: return "reject"
    if "major" in vl:  return "major"
    if "minor" in vl:  return "minor"
    return "accept"


def compute_agreement_stats() -> dict:
    log = st.session_state.agreement_log
    if not log:
        return {"total": 0, "agree": 0, "partial": 0, "disagree": 0, "rate": None}
    total   = len(log)
    agree   = sum(1 for r in log if r["feedback"] == "agree")
    partial = sum(1 for r in log if r["feedback"] == "partial")
    disagree= sum(1 for r in log if r["feedback"] == "disagree")
    rate    = round((agree + 0.5 * partial) / total * 100, 1) if total else None
    return {"total": total, "agree": agree, "partial": partial, "disagree": disagree, "rate": rate}


# ─── SIDEBAR ──────────────────────────────────────────────────────────────────
with st.sidebar:
    st.markdown(
        """
        <div style="padding:0.5rem 0 1rem;">
          <div style="font-family:'Playfair Display',Georgia,serif;
                      font-size:1.4rem;font-weight:600;color:#f0ebe0;
                      letter-spacing:-0.01em;line-height:1.2;">
            DentEdTech™
          </div>
          <div style="font-family:'Source Serif 4',Georgia,serif;
                      font-size:0.72rem;font-weight:300;font-style:italic;
                      color:#7a8a9a;margin-top:0.2rem;letter-spacing:0.04em;">
            Manuscript Intelligence Platform
          </div>
          <div style="width:24px;height:1px;background:#b5903a;margin-top:0.8rem;"></div>
        </div>
        """,
        unsafe_allow_html=True,
    )
    st.caption(f"Partner: {st.session_state.access_partner}")

    # Open access — no cap display

    with st.expander("🔒 Data & Privacy Status", expanded=False):
        st.markdown(
            f"""
            | Item | Status |
            |---|---|
            | Access granted | ✅ {st.session_state.access_partner} |
            | Files written to disk | ✅ Never |
            | API provider | Anthropic |
            | Training on API data | ✅ No |
            | Session started | {st.session_state.session_start.strftime('%H:%M UTC')} |
            | Documents processed | {st.session_state.upload_count} |
            """
        )
        if st.session_state.pdf_hash:
            st.caption(f"SHA-256: `{st.session_state.pdf_hash[:16]}…`")
        st.markdown("[Privacy Policy](https://www.anthropic.com/privacy) · [Request ZDR](https://www.anthropic.com/contact-sales)")

    # Agreement tracking panel
    stats = compute_agreement_stats()
    if stats["total"] > 0:
        with st.expander(f"📊 Agreement Tracking ({stats['total']} reviews)", expanded=False):
            if stats["rate"] is not None:
                st.metric("Agreement rate", f"{stats['rate']}%")
            c1, c2, c3 = st.columns(3)
            c1.metric("✅ Agree",    stats["agree"])
            c2.metric("⚠️ Partial", stats["partial"])
            c3.metric("❌ Disagree", stats["disagree"])
            st.caption("Anonymous — no manuscript data stored.")

    st.divider()

    uploaded = st.file_uploader("Upload manuscript (PDF)", type=["pdf"])
    if uploaded:
        if uploaded.name != st.session_state.pdf_name or not st.session_state.pdf_base64:
            with st.spinner("Encoding PDF in memory…"):
                b64, txt, sha = encode_pdf(uploaded)
            if sha != st.session_state.pdf_hash:
                st.session_state.pdf_base64        = b64
                st.session_state.pdf_text          = txt
                st.session_state.pdf_hash          = sha
                st.session_state.pdf_name          = uploaded.name
                st.session_state.report            = None
                st.session_state.raw_report        = ""
                st.session_state.chat_history      = []
                st.session_state.similarity_report = None
                st.session_state.raw_similarity    = ""
                st.session_state.search_results    = []
                st.session_state.feedback_given    = False
                st.session_state.upload_count     += 1
                st.success(f"✅ New file: {uploaded.name}")
            else:
                st.success(f"✅ {uploaded.name} (unchanged)")
        else:
            st.success(f"✅ {uploaded.name}")
        st.caption(f"{len(st.session_state.pdf_text):,} chars · SHA-256: {st.session_state.pdf_hash[:12]}…")

    st.divider()
    journal = st.selectbox("Target journal", JOURNALS)

    st.markdown("**Study design**")
    study_design = st.selectbox(
        "Study design",
        STUDY_DESIGNS,
        label_visibility="collapsed",
        help="Select the study design to apply appropriate criteria. 'Auto-detect' lets the AI identify it.",
    )

    st.markdown("**Review criteria**")
    selected_criteria = [
        k for k, label in REVIEW_CRITERIA.items()
        if st.checkbox(label, value=True, key=f"cb_{k}")
    ]

    st.divider()
    can_analyze = (
        bool(st.session_state.pdf_base64)
        and len(selected_criteria) > 0
        and check_session_cap()
    )
    if st.button("🚀 Run Full Analysis", disabled=not can_analyze, use_container_width=True):
        st.session_state.report         = None
        st.session_state.raw_report     = ""
        st.session_state.chat_history   = []
        st.session_state.feedback_given = False
        st.session_state["_trigger_analysis"] = True

    st.divider()
    if st.button("🗑️ Clear Session Data", use_container_width=True):
        clear_session_data()
        st.success("Session cleared.")
        st.rerun()

    st.caption("⚠️ All data clears when you close this tab.")

# ─── ANALYSIS ─────────────────────────────────────────────────────────────────
if st.session_state.get("_trigger_analysis"):
    st.session_state["_trigger_analysis"] = False

    if not check_session_cap():
        st.error("Session analysis limit reached. Contact DentEdTech™ for extended access.")
        st.stop()

    system = build_system_prompt(journal)
    prompt = build_review_prompt(selected_criteria, journal, study_design)
    phases = [
        "Phase 1 — Deep document read & structure audit",
        "Phase 2 — Methodology & criteria assessment",
        "Phase 3 — Citation audit & gap analysis",
        "Phase 4 — Generating structured review report",
    ]

    progress = st.progress(0)
    status   = st.status("Running analysis…", expanded=True)
    raw = None; model_used = PRIMARY_MODEL

    for i, phase in enumerate(phases):
        status.write(f"⚙️ {phase}")
        progress.progress((i + 1) / len(phases))

    try:
        status.write(f"🧠 Sending to {PRIMARY_MODEL}…")
        raw = call_api_with_pdf(system, prompt, PRIMARY_MODEL, max_tok=6000)
        model_used = PRIMARY_MODEL
    except NotFoundError:
        status.write(f"⚠️ Falling back to {FALLBACK_MODEL}…")
        try:
            raw = call_api_with_pdf(system, prompt, FALLBACK_MODEL, max_tok=6000)
            model_used = FALLBACK_MODEL
        except Exception:
            raw = call_api_with_text(system, prompt, FALLBACK_MODEL, max_tok=6000)
            model_used = FALLBACK_MODEL + " (text)"
    except Exception as e:
        try:
            raw = call_api_with_text(system, prompt, PRIMARY_MODEL, max_tok=6000)
            model_used = PRIMARY_MODEL + " (text)"
        except Exception as e2:
            status.update(label=f"Error: {e2}", state="error")
            st.stop()

    progress.progress(1.0)
    st.session_state.report     = parse_json(raw)
    st.session_state.raw_report = raw
    st.session_state.model_used = model_used
    st.session_state.analyses_this_session += 1

    # Build a plain-English chat seed so the editor chat responds in prose not JSON
    parsed_seed = parse_json(raw)
    if parsed_seed:
        verdict    = parsed_seed.get("verdict", "—")
        score      = parsed_seed.get("overall_score", "—")
        confidence = parsed_seed.get("confidence", "—")
        summary    = parsed_seed.get("executive_summary", "")
        kp         = parsed_seed.get("kirkpatrick_level", {})
        kp_text    = f"Kirkpatrick Level {kp.get('level','?')}: {kp.get('justification','')}" if kp else "—"
        strengths  = parsed_seed.get("strengths", [])
        weaknesses = parsed_seed.get("weaknesses", [])
        recs       = parsed_seed.get("actionable_recommendations", [])
        seed_msg = (
            f"I have completed a full structured peer review of this manuscript.\n\n"
            f"VERDICT: {verdict} | Score: {score}/100 | Confidence: {confidence}\n\n"
            f"SUMMARY: {summary}\n\n"
            f"KIRKPATRICK: {kp_text}\n\n"
            f"STRENGTHS ({len(strengths)}):\n" +
            "\n".join(f"- {s}" for s in strengths[:5]) + "\n\n"
            f"KEY WEAKNESSES ({len(weaknesses)}):\n" +
            "\n".join(f"- [{w.get('severity','').upper()}] {w.get('section','')}: {w.get('issue','')}" for w in weaknesses[:5]) + "\n\n"
            f"TOP RECOMMENDATIONS:\n" +
            "\n".join(f"{i+1}. {r}" for i, r in enumerate(recs[:5])) + "\n\n"
            f"I am ready to answer any questions about this review or the manuscript in detail."
        )
    else:
        seed_msg = (
            "I have completed a full structured peer review of this manuscript. "
            "I am ready to answer any questions about the methodology, citations, "
            "structure, or any aspect of the review in plain language."
        )

    st.session_state.chat_history = [
        {"role": "user", "content": [
            {"type": "document", "source": {"type": "base64", "media_type": "application/pdf",
                                            "data": st.session_state.pdf_base64}},
            {"type": "text", "text": "This is the manuscript we just reviewed. Please answer all my questions about it in clear, plain English — never return JSON."},
        ]},
        {"role": "assistant", "content": seed_msg},
    ]
    status.update(label="Analysis complete ✓", state="complete", expanded=False)
    st.rerun()

# ─── IDLE STATE ───────────────────────────────────────────────────────────────
if not st.session_state.report and not st.session_state.raw_report:
    st.markdown(
        """
        <div style="max-width:680px;margin:5rem auto;text-align:center;">
          <div style="font-family:'Source Serif 4',Georgia,serif;
                      font-size:0.68rem;letter-spacing:0.32em;text-transform:uppercase;
                      color:#8a9aaa;margin-bottom:1.2rem;">
            Health Professions Education · Manuscript Intelligence
          </div>
          <h1 style="font-family:'Playfair Display',Georgia,serif;
                     font-size:3rem;font-weight:600;color:#0f2535;
                     margin:0 0 0.4rem;letter-spacing:-0.03em;line-height:1.1;">
            DentEdTech™
          </h1>
          <div style="font-family:'Source Serif 4',Georgia,serif;
                      font-size:1.05rem;font-weight:300;font-style:italic;
                      color:#6a7a8a;margin-bottom:2rem;">
            Expert peer review intelligence for dental &amp; health professions education journals
          </div>
          <div style="width:48px;height:1px;background:#b5903a;margin:0 auto 2.5rem;"></div>
          <p style="font-family:'Source Serif 4',Georgia,serif;
                    font-size:0.92rem;color:#6a6460;line-height:1.85;
                    max-width:480px;margin:0 auto 3rem;">
            Upload a manuscript PDF using the sidebar, select your target journal
            and review criteria, then run the analysis.
          </p>
          <div style="display:flex;flex-wrap:wrap;gap:8px;justify-content:center;">
            <span style="font-family:'Source Serif 4',serif;font-size:0.74rem;
                         letter-spacing:0.06em;text-transform:uppercase;
                         color:#1a3a4a;background:#f0ebe0;border:1px solid #ddd8cc;
                         padding:5px 14px;border-radius:2px;">Confidence scoring</span>
            <span style="font-family:'Source Serif 4',serif;font-size:0.74rem;
                         letter-spacing:0.06em;text-transform:uppercase;
                         color:#1a3a4a;background:#f0ebe0;border:1px solid #ddd8cc;
                         padding:5px 14px;border-radius:2px;">CONSORT · SRQR · COREQ</span>
            <span style="font-family:'Source Serif 4',serif;font-size:0.74rem;
                         letter-spacing:0.06em;text-transform:uppercase;
                         color:#1a3a4a;background:#f0ebe0;border:1px solid #ddd8cc;
                         padding:5px 14px;border-radius:2px;">Kirkpatrick framework</span>
            <span style="font-family:'Source Serif 4',serif;font-size:0.74rem;
                         letter-spacing:0.06em;text-transform:uppercase;
                         color:#1a3a4a;background:#f0ebe0;border:1px solid #ddd8cc;
                         padding:5px 14px;border-radius:2px;">Citation audit</span>
            <span style="font-family:'Source Serif 4',serif;font-size:0.74rem;
                         letter-spacing:0.06em;text-transform:uppercase;
                         color:#1a3a4a;background:#f0ebe0;border:1px solid #ddd8cc;
                         padding:5px 14px;border-radius:2px;">Web similarity search</span>
            <span style="font-family:'Source Serif 4',serif;font-size:0.74rem;
                         letter-spacing:0.06em;text-transform:uppercase;
                         color:#1a3a4a;background:#f0ebe0;border:1px solid #ddd8cc;
                         padding:5px 14px;border-radius:2px;">Author feedback report</span>
            <span style="font-family:'Source Serif 4',serif;font-size:0.74rem;
                         letter-spacing:0.06em;text-transform:uppercase;
                         color:#1a3a4a;background:#f0ebe0;border:1px solid #ddd8cc;
                         padding:5px 14px;border-radius:2px;">Agreement tracking</span>
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )
    st.stop()

# ─── MAIN DISPLAY ─────────────────────────────────────────────────────────────
report = st.session_state.report
raw    = st.session_state.raw_report

show_cap_warning()
st.markdown(
    f"<div style='font-family:Source Serif 4,serif;font-size:0.75rem;"
    f"font-style:italic;color:#8a847a;margin-bottom:0.5rem;'>"
    f"Analysis generated with {st.session_state.model_used}</div>",
    unsafe_allow_html=True,
)

tab_report, tab_author, tab_similarity, tab_chat, tab_feedback, tab_humanise = st.tabs([
    "Editor Report",
    "Author Feedback",
    "Similarity Audit",
    "Editor Chat",
    "Feedback",
    "Humanise Writing",
])

# ─── REPORT TAB ───────────────────────────────────────────────────────────────
with tab_report:
    if report is None:
        st.warning("Could not parse structured JSON — showing raw report.")
        st.text_area("Raw report", raw, height=600)
    else:
        verdict    = report.get("verdict", "Unknown")
        score      = report.get("overall_score", "—")
        confidence = report.get("confidence", "Unknown")
        conf_note  = report.get("confidence_note", "")
        detected_design = report.get("detected_study_design", "")

        col_v, col_c, col_s, col_dl = st.columns([3, 2, 1, 1])
        with col_v:
            st.markdown(render_badge(verdict, verdict_key(verdict), VERDICT_COLOURS), unsafe_allow_html=True)
            if detected_design:
                st.caption(f"Study design: {detected_design}")
        with col_c:
            st.markdown(render_badge(f"Confidence: {confidence}", confidence.lower(), CONFIDENCE_COLOURS), unsafe_allow_html=True)
        with col_s:
            st.metric("Score", f"{score}/100")
        with col_dl:
            st.download_button(
                "⬇️ Editor .docx",
                data=create_docx(report, raw),
                file_name="DentEdTech_Editor_Report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

        if confidence.lower() in ("moderate", "low"):
            st.warning(
                f"⚠️ **Confidence flag — {confidence}:** {conf_note}\n\n"
                "*The AI has flagged uncertainty about this verdict. "
                "Editor verification is strongly recommended before any desk-rejection decision.*"
            )
        elif conf_note:
            st.info(f"ℹ️ **Confidence note:** {conf_note}")

        st.markdown(f"> {report.get('executive_summary','')}")
        st.divider()

        scores = report.get("scores", {})
        if scores:
            cols = st.columns(len(scores))
            for col, (k, v) in zip(cols, scores.items()):
                col.metric(k.capitalize(), f"{v}/10")

        kp = report.get("kirkpatrick_level", {})
        if kp:
            st.info(f"🎯 **Kirkpatrick Level {kp.get('level','?')}** — {kp.get('justification','')}")

        st.divider()

        with st.expander("🧵 Golden Thread Analysis", expanded=True):
            st.write(report.get("golden_thread", "—"))

        if report.get("strengths"):
            with st.expander(f"✅ Strengths ({len(report['strengths'])})", expanded=True):
                for s in report["strengths"]:
                    st.markdown(f"- {s}")

        weaknesses = report.get("weaknesses", [])
        if weaknesses:
            majors = [w for w in weaknesses if w.get("severity") == "major"]
            minors = [w for w in weaknesses if w.get("severity") != "major"]
            with st.expander(f"⚠️ Weaknesses — {len(majors)} major, {len(minors)} minor", expanded=True):
                for w in weaknesses:
                    sev = w.get("severity","minor")
                    st.markdown(
                        f"{'🔴' if sev=='major' else '🟡'} "
                        f"**[{sev.upper()} — {w.get('section','')}]** {w.get('issue','')}"
                    )
                    if w.get("suggestion"):
                        st.caption(f"→ {w['suggestion']}")

        sc = report.get("section_comments", {})
        if sc:
            with st.expander("📝 Section-by-Section Comments"):
                for section, comment in sc.items():
                    st.markdown(f"**{section.capitalize()}**")
                    st.write(comment)
                    st.divider()

        ca = report.get("citation_audit", {})
        if ca:
            with st.expander("📚 Citation Audit"):
                for ref in ca.get("missing_key_references", []):
                    st.markdown(f"- Missing: {ref}")
                for ref in ca.get("potentially_outdated", []):
                    st.markdown(f"- Outdated: {ref}")
                st.markdown(f"**Mismatches:** {ca.get('mismatches','None identified')}")

        recs = report.get("actionable_recommendations", [])
        if recs:
            with st.expander(f"✅ Actionable Recommendations ({len(recs)})", expanded=True):
                for i, rec in enumerate(recs, 1):
                    st.markdown(f"**{i}.** {rec}")

        # ── Design-specific assessment ─────────────────────────────────────────
        dsa = report.get("design_specific_assessment", {})
        if dsa:
            with st.expander("🔬 Study Design — Specific Assessment", expanded=True):
                na_fields = []
                active_fields = []
                for k, v in dsa.items():
                    label = k.replace("_", " ").title()
                    v_str = str(v)
                    if "N/A" in v_str or "n/a" in v_str.lower() or "not applicable" in v_str.lower():
                        na_fields.append(label)
                    else:
                        active_fields.append((label, v_str))
                for label, val in active_fields:
                    st.markdown(f"**{label}:** {val}")
                if na_fields:
                    st.caption(f"Not applicable for this study design: {', '.join(na_fields)}")

        # ── Writing quality assessment ─────────────────────────────────────────
        wqa = report.get("writing_quality_assessment", {})
        if wqa:
            scores = report.get("scores", {})
            wq_score = scores.get("writing_quality", "—")
            hum_score = scores.get("humanisation", "—")
            with st.expander(f"✍️ Writing Quality & Humanisation — {wq_score}/10 writing · {hum_score}/10 humanisation", expanded=True):
                if wqa.get("academic_register"):
                    st.markdown(f"**Academic register:** {wqa['academic_register']}")
                if wqa.get("voice_consistency"):
                    st.markdown(f"**Voice consistency:** {wqa['voice_consistency']}")
                if wqa.get("humanisation_score"):
                    st.markdown(f"**Humanisation:** {wqa['humanisation_score']}")
                generic = wqa.get("generic_phrases_found", [])
                if generic:
                    st.markdown("**Generic / hollow phrases to replace:**")
                    for phrase in generic:
                        st.warning(f'🔴 "{phrase}"')
                suggestions = wqa.get("humanisation_suggestions", [])
                if suggestions:
                    st.markdown("**Rewrite suggestions:**")
                    for i, sug in enumerate(suggestions, 1):
                        st.success(f"✏️ {i}. {sug}")

        if report.get("editor_note"):
            with st.expander("🔒 Confidential Note to Editor"):
                st.info(report["editor_note"])

        # Agreement feedback
        st.divider()
        st.markdown("#### 📊 Editor Agreement Feedback")
        st.caption("Does this AI verdict match your own assessment? Anonymous feedback only.")

        if not st.session_state.feedback_given:
            col_a, col_b, col_c = st.columns(3)
            for label, val, col in [
                ("✅ I agree", "agree", col_a),
                ("⚠️ Partially agree", "partial", col_b),
                ("❌ I disagree", "disagree", col_c),
            ]:
                if col.button(label, use_container_width=True):
                    st.session_state.agreement_log.append({
                        "verdict": verdict, "score": score,
                        "feedback": val,
                        "timestamp": datetime.datetime.utcnow().isoformat(),
                    })
                    st.session_state.feedback_given = True
                    st.rerun()
        else:
            last = st.session_state.agreement_log[-1]["feedback"]
            icons = {"agree": "✅", "partial": "⚠️", "disagree": "❌"}
            st.success(f"{icons.get(last,'✅')} Feedback recorded — thank you.")
            stats = compute_agreement_stats()
            if stats["total"] >= 3:
                st.info(f"📊 Running agreement rate across {stats['total']} reviews: **{stats['rate']}%**")

# ─── AUTHOR FEEDBACK TAB ─────────────────────────────────────────────────────
with tab_author:
    st.info(
        "This tab generates a **constructive author-facing report** — "
        "no verdict, no editor note. Safe to share directly with authors."
    )
    if report is None:
        st.warning("Run a full analysis first.")
    else:
        with st.expander("👁️ Preview author feedback", expanded=True):
            st.markdown("**Overall feedback:**")
            st.write(report.get("author_feedback_summary", "No summary available."))
            st.markdown("**Strengths:**")
            for s in report.get("strengths", []):
                st.markdown(f"- {s}")
            st.markdown("**Priority actions before resubmission:**")
            for i, rec in enumerate(report.get("actionable_recommendations", []), 1):
                st.markdown(f"**{i}.** {rec}")

        st.download_button(
            "⬇️ Download Author Feedback Report (.docx)",
            data=create_author_feedback_docx(report),
            file_name="DentEdTech_Author_Feedback.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
        st.caption("This report contains no editorial verdict or confidential content.")

# ─── SIMILARITY TAB ───────────────────────────────────────────────────────────
with tab_similarity:
    st.info(
        "⚠️ **Disclaimer:** Searches open-access content only. "
        "Not equivalent to Turnitin or iThenticate. "
        "Always use your institution's official similarity checker before submission."
    )

    if not DDG_AVAILABLE:
        st.warning("Add `duckduckgo-search` to `requirements.txt` to enable web search.")

    if st.session_state.pdf_name:
        col_f, col_clr = st.columns([4, 1])
        with col_f:
            st.caption(f"📄 **{st.session_state.pdf_name}** · SHA-256: `{st.session_state.pdf_hash[:16]}…`")
        with col_clr:
            if st.button("🔄 Clear audit"):
                st.session_state.similarity_report = None
                st.session_state.raw_similarity    = ""
                st.session_state.search_results    = []
                st.rerun()

    if st.button("🔍 Run Similarity & Originality Audit",
                 disabled=not bool(st.session_state.pdf_base64)):
        st.session_state.similarity_report = None
        st.session_state.raw_similarity    = ""
        st.session_state.search_results    = []

        with st.status("Running similarity audit…", expanded=True) as ss:
            ss.write("🔎 Step 1 — Extracting key phrases…")
            queries = extract_search_queries(st.session_state.pdf_text)
            ss.write(f"   {len(queries)} queries extracted")

            if DDG_AVAILABLE:
                ss.write("🌐 Step 2 — Searching open-access publications…")
                sr = search_web(queries)
                st.session_state.search_results = sr
                ss.write(f"   {len(sr)} papers found")
            else:
                sr = []
                ss.write("⚠️ Step 2 — Web search skipped")

            ss.write("🧠 Step 3 — AI originality analysis…")
            sim_prompt  = build_similarity_prompt(sr)
            sim_system  = (
                "You are an academic integrity specialist. Analyse manuscripts for "
                "similarity risks, boilerplate, paraphrase patterns, and published overlap. "
                "Be precise, quote passages, never fabricate."
            )
            sim_raw = None
            try:
                sim_raw = call_api_with_pdf(sim_system, sim_prompt, FALLBACK_MODEL, max_tok=8000)
            except Exception as e:
                ss.write(f"⚠️ PDF mode failed ({e}) — text mode…")
                try:
                    sim_raw = call_api_with_text(sim_system, sim_prompt, FALLBACK_MODEL, max_tok=8000)
                except Exception as e2:
                    ss.update(label=f"Error: {e2}", state="error"); st.stop()

            parsed_sim = parse_json(sim_raw)
            st.session_state.similarity_report = parsed_sim
            st.session_state.raw_similarity    = sim_raw
            ss.update(label="Similarity audit complete ✓", state="complete", expanded=False)
        st.rerun()

    sim = st.session_state.similarity_report
    if sim is None and st.session_state.raw_similarity:
        sim = parse_json(st.session_state.raw_similarity)
        if sim:
            st.session_state.similarity_report = sim
        else:
            # Parse totally failed — render a readable plain-text report from the raw response
            raw_d = st.session_state.raw_similarity.replace("```json","").replace("```","").strip()
            st.info(
                "The AI returned a response that could not be fully structured. "
                "The full analysis is displayed below as plain text."
            )
            # Try to extract key fields manually and show them nicely
            import re
            def extract_field(text, key):
                pattern = rf'"{key}"\s*:\s*"([^"]+)"'
                m = re.search(pattern, text)
                if m: return m.group(1)
                pattern2 = rf'"{key}"\s*:\s*(\d+)'
                m2 = re.search(pattern2, text)
                if m2: return m2.group(1)
                return None

            risk   = extract_field(raw_d, "overall_risk_level") or "See full response below"
            est    = extract_field(raw_d, "estimated_similarity_risk_percent") or "—"
            ready  = extract_field(raw_d, "submission_readiness") or ""
            disc   = extract_field(raw_d, "disclaimer") or ""

            if risk != "See full response below":
                col1, col2 = st.columns([3, 1])
                with col1:
                    st.markdown(render_badge(f"Risk: {risk}", risk.lower(), RISK_COLOURS), unsafe_allow_html=True)
                with col2:
                    st.metric("Estimated %", f"~{est}%")
                if disc:
                    st.caption(disc)
                if ready:
                    st.markdown(f"**Submission readiness:** {ready}")
                st.divider()

            # Show full response as readable text, not JSON
            # Strip JSON structure markers and show as prose
            clean = re.sub(r'[\{\}\[\]"]', " ", raw_d)
            clean = re.sub(r',\s*\n', "\n", clean)
            clean = re.sub(r'\s{2,}', " ", clean)
            lines = [l.strip() for l in clean.split("\n") if l.strip() and len(l.strip()) > 3]
            for line in lines:
                if ":" in line:
                    parts = line.split(":", 1)
                    key = parts[0].strip().replace("_", " ").title()
                    val = parts[1].strip() if len(parts) > 1 else ""
                    if val and len(val) > 2:
                        st.markdown(f"**{key}:** {val}")
                elif len(line) > 10:
                    st.markdown(f"- {line}")

    elif sim:
        risk = sim.get("overall_risk_level", "Unknown")
        est  = sim.get("estimated_similarity_risk_percent", "—")
        col1, col2 = st.columns([3, 1])
        with col1:
            st.markdown(render_badge(f"Risk: {risk}", risk.lower(), RISK_COLOURS), unsafe_allow_html=True)
        with col2:
            st.metric("Estimated %", f"~{est}%", help="AI estimate — not Turnitin")
        st.caption(sim.get("disclaimer",""))
        st.markdown(f"**Submission readiness:** {sim.get('submission_readiness','—')}")
        st.divider()

        pubs = sim.get("similar_publications", [])
        label = f"🌐 Similar Publications Found Online ({len(pubs)})" if pubs else "🌐 Similar Publications Found Online"
        with st.expander(label, expanded=bool(pubs)):
            if pubs:
                for p in pubs:
                    rl = p.get("risk_level","")
                    st.markdown(
                        f"**{p.get('title','')}** "
                        f"{render_badge(rl, rl.lower(), RISK_COLOURS)}",
                        unsafe_allow_html=True,
                    )
                    if p.get("url"): st.markdown(f"🔗 [{p['url']}]({p['url']})")
                    st.markdown(f"**Overlap:** {p.get('overlap_description','')}")
                    if p.get("recommendation"): st.success(f"✏️ {p['recommendation']}")
                    st.divider()
            else:
                st.success("✅ No significantly similar papers identified.")

        if st.session_state.search_results:
            with st.expander(f"🔎 Raw Search Results ({len(st.session_state.search_results)})", expanded=False):
                for r in st.session_state.search_results:
                    st.markdown(f"**{r['title']}**")
                    st.caption(f"`{r['query']}` · {r['url']}")
                    st.write(r["body"][:200] + "…")
                    st.divider()

        boiler = sim.get("boilerplate_sections", [])
        with st.expander(f"📋 Boilerplate Sections ({len(boiler)})" if boiler else "📋 Boilerplate Sections", expanded=bool(boiler)):
            if boiler:
                for item in boiler:
                    st.markdown(f"**Section:** {item.get('section','')}")
                    st.warning(f"🔴 **High-risk passage:**\n> {item.get('passage','')}")
                    st.caption(f"Risk: {item.get('risk','')}")
                    st.success(f"✏️ {item.get('suggestion','')}")
                    st.divider()
            else:
                st.success("✅ No significant boilerplate detected.")

        claims = sim.get("citation_free_claims", [])
        with st.expander(f"📎 Citation-Free Claims ({len(claims)})" if claims else "📎 Citation-Free Claims", expanded=bool(claims)):
            if claims:
                for item in claims:
                    st.warning(f"🟡 > {item.get('passage','')}")
                    st.caption(item.get('risk',''))
                    st.success(f"✏️ {item.get('suggestion','')}")
                    st.divider()
            else:
                st.success("✅ No significant uncited claims detected.")

        repeats = sim.get("internal_repetition", [])
        with st.expander(f"🔁 Internal Repetition ({len(repeats)})" if repeats else "🔁 Internal Repetition", expanded=bool(repeats)):
            if repeats:
                for item in repeats:
                    st.warning(f"🟡 > {item.get('passage','')}")
                    st.caption(f"Appears in: {', '.join(item.get('appears_in',[]))}")
                    st.divider()
            else:
                st.success("✅ No significant internal repetition detected.")

        if sim.get("methods_risk"):
            with st.expander("🔬 Methods Section Originality", expanded=True):
                st.write(sim["methods_risk"])

        rewrites = sim.get("priority_rewrites", [])
        if rewrites:
            with st.expander(f"✏️ Priority Rewrites ({len(rewrites)})", expanded=True):
                for i, rw in enumerate(rewrites, 1):
                    st.markdown(f"**{i}.** {rw}")

# ─── CHAT TAB ─────────────────────────────────────────────────────────────────
with tab_chat:
    st.caption("Ask questions about the review or manuscript. Full PDF is in context.")

    quick_prompts = [
        "Expand on the methodology critique",
        "Which specific citations are missing and why?",
        "How can the Discussion be strengthened?",
        "Explain the golden thread score in detail",
        "What would reach Kirkpatrick Level 3 or 4?",
        "Suggest a revised abstract",
    ]
    cols = st.columns(3)
    for i, qp in enumerate(quick_prompts):
        if cols[i % 3].button(qp, key=f"qp_{i}", use_container_width=True):
            st.session_state._pending_chat = qp

    st.divider()

    for msg in st.session_state.chat_history[2:]:
        role    = msg["role"]
        content = msg["content"] if isinstance(msg["content"], str) else str(msg["content"])
        with st.chat_message(role):
            st.markdown(content)

    pending    = st.session_state.pop("_pending_chat", None)
    user_input = st.chat_input("Ask about the review or manuscript…") or pending

    if user_input:
        with st.chat_message("user"):
            st.markdown(user_input)
        st.session_state.chat_history.append({"role": "user", "content": user_input})

        with st.chat_message("assistant"):
            with st.spinner("Thinking…"):
                try:
                    response = client.messages.create(
                        model=FALLBACK_MODEL,
                        max_tokens=2048,
                        system=(
                            "You are a Senior HPE Journal Editor discussing a peer review you just completed. "
                            "Answer every question in clear, natural, conversational English prose. "
                            "Never return JSON, never use code blocks, never use structured data formats. "
                            "Quote specific manuscript passages when relevant. "
                            "Be constructive, precise, and suggest concrete improvements. "
                            "Write as an expert colleague speaking directly to another editor."
                        ),
                        messages=st.session_state.chat_history,
                    )
                    reply = response.content[0].text
                except Exception as e:
                    reply = f"Error: {e}"
            st.markdown(reply)
        st.session_state.chat_history.append({"role": "assistant", "content": reply})
# ─── FEEDBACK TAB ─────────────────────────────────────────────────────────────
with tab_feedback:
    st.markdown(
        """
        <div style="max-width:640px;margin:0 auto;">
          <div style="font-family:'Playfair Display',Georgia,serif;font-size:1.6rem;
                      font-weight:600;color:#0f2535;margin-bottom:0.3rem;">
            Platform Feedback
          </div>
          <div style="font-family:'Source Serif 4',Georgia,serif;font-size:0.9rem;
                      font-style:italic;color:#6a7a8a;margin-bottom:1.5rem;">
            Help us improve DentEdTech™ — your feedback is sent directly to the development team.
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    if st.session_state.feedback_submitted:
        st.success(
            "✅ Thank you — your feedback has been recorded. "
            "Please click the button below to send it directly to the DentEdTech™ team."
        )
        col_a, col_b = st.columns(2)
        with col_a:
            mailto = st.session_state.get("_last_mailto", "mailto:ayman.khalifah@manchester.ac.uk")
            st.markdown(
                f'<a href="{mailto}" target="_blank">'
                f'<button style="width:100%;padding:0.55rem 1rem;'
                f'background:#0f2535;color:#f0ebe0;border:none;border-radius:3px;'
                f'font-family:Source Serif 4,serif;font-size:0.88rem;'
                f'letter-spacing:0.04em;cursor:pointer;">'
                f'📧 Send Feedback Email</button></a>',
                unsafe_allow_html=True,
            )
        with col_b:
            csv_data = st.session_state.get("_last_feedback_csv", b"")
            if csv_data:
                st.download_button(
                    "⬇️ Download as CSV",
                    data=csv_data,
                    file_name="DentEdTech_Feedback.csv",
                    mime="text/csv",
                    use_container_width=True,
                )
        st.caption(
            "The email button opens your mail client pre-filled with your feedback. "
            "The CSV lets you save a copy for your records."
        )
        st.divider()
        if st.button("Submit another response"):
            st.session_state.feedback_submitted = False
            st.rerun()
    else:
        with st.form("feedback_form", clear_on_submit=True):
            st.markdown("#### About you")
            col_name, col_role = st.columns(2)
            with col_name:
                fb_name = st.text_input(
                    "Your name (optional)",
                    placeholder="e.g. Dr. Jane Smith",
                )
            with col_role:
                fb_role = st.selectbox(
                    "Your role",
                    [
                        "Journal Editor",
                        "Managing Editor",
                        "Academic Researcher",
                        "Dental Educator",
                        "HPE Educator",
                        "Publishing Professional",
                        "Other",
                    ],
                )
            fb_institution = st.text_input(
                "Institution / Journal (optional)",
                placeholder="e.g. European Journal of Dental Education",
            )
            fb_email = st.text_input(
                "Your email (optional — only if you would like a response)",
                placeholder="e.g. j.smith@university.ac.uk",
            )

            st.divider()
            st.markdown("#### Your experience")

            fb_overall = st.select_slider(
                "Overall satisfaction with DentEdTech™",
                options=["Very dissatisfied", "Dissatisfied", "Neutral", "Satisfied", "Very satisfied"],
                value="Satisfied",
            )

            fb_accuracy = st.select_slider(
                "How accurate was the AI review compared to your own editorial assessment?",
                options=["Very inaccurate", "Somewhat inaccurate", "Neutral", "Somewhat accurate", "Very accurate"],
                value="Somewhat accurate",
            )

            fb_usefulness = st.select_slider(
                "How useful would this tool be in your editorial workflow?",
                options=["Not useful", "Slightly useful", "Moderately useful", "Very useful", "Essential"],
                value="Moderately useful",
            )

            st.divider()
            st.markdown("#### Detailed feedback")

            fb_worked_well = st.text_area(
                "What worked well? What did you find most valuable?",
                placeholder="e.g. The confidence scoring helped me understand when to apply extra scrutiny...",
                height=100,
            )

            fb_improvements = st.text_area(
                "What needs improvement? What frustrated you or felt inaccurate?",
                placeholder="e.g. The citation audit missed some key recent papers...",
                height=100,
            )

            fb_features = st.text_area(
                "What features or capabilities would you like to see added?",
                placeholder="e.g. Integration with submission systems, batch processing, language support...",
                height=100,
            )

            fb_recommend = st.radio(
                "Would you recommend DentEdTech™ to a colleague?",
                ["Definitely yes", "Probably yes", "Not sure", "Probably not", "Definitely not"],
                horizontal=True,
            )

            fb_pilot = st.radio(
                "Would you be interested in a formal pilot partnership with your journal?",
                ["Yes, very interested", "Possibly interested", "Not at this stage", "Not applicable"],
                horizontal=True,
            )

            fb_other = st.text_area(
                "Any other comments, suggestions, or questions?",
                placeholder="Anything else you would like us to know...",
                height=80,
            )

            st.divider()
            submitted = st.form_submit_button(
                "Submit Feedback →",
                use_container_width=True,
            )

        if submitted:
            import urllib.parse
            import csv
            import io
            import datetime as dt

            # Build feedback summary
            feedback_data = {
                "Timestamp":            dt.datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC"),
                "Partner":              st.session_state.access_partner,
                "Name":                 fb_name or "Anonymous",
                "Role":                 fb_role,
                "Institution":          fb_institution or "Not provided",
                "Email":                fb_email or "Not provided",
                "Overall satisfaction": fb_overall,
                "Accuracy":             fb_accuracy,
                "Usefulness":           fb_usefulness,
                "What worked well":     fb_worked_well or "Not provided",
                "Improvements needed":  fb_improvements or "Not provided",
                "Feature requests":     fb_features or "Not provided",
                "Would recommend":      fb_recommend,
                "Pilot interest":       fb_pilot,
                "Other comments":       fb_other or "None",
            }

            # Store in session for display
            if "all_feedback" not in st.session_state:
                st.session_state.all_feedback = []
            st.session_state.all_feedback.append(feedback_data)

            # Build mailto link so user can send directly
            body_lines = [f"{k}: {v}" for k, v in feedback_data.items()]
            body_text  = "%0A".join(body_lines)
            subject    = urllib.parse.quote(
                f"DentEdTech Feedback — {fb_role} — {fb_overall}"
            )
            mailto_url = (
                f"mailto:ayman.khalifah@manchester.ac.uk"
                f"?subject={subject}&body={body_text}"
            )

            # Build CSV for download
            csv_buf = io.StringIO()
            writer  = csv.DictWriter(csv_buf, fieldnames=list(feedback_data.keys()))
            writer.writeheader()
            writer.writerow(feedback_data)
            csv_bytes = csv_buf.getvalue().encode("utf-8")

            st.session_state.feedback_submitted   = True
            st.session_state._last_mailto         = mailto_url
            st.session_state._last_feedback_csv   = csv_bytes
            st.rerun()


# ─── HUMANISE WRITING TAB ─────────────────────────────────────────────────────
with tab_humanise:
    st.markdown(
        """
        <div style="font-family:'Playfair Display',Georgia,serif;font-size:1.3rem;
                    font-weight:600;color:#0f2535;margin-bottom:0.3rem;">
            Humanise Your Writing
        </div>
        <div style="font-family:'Source Serif 4',Georgia,serif;font-size:0.88rem;
                    font-style:italic;color:#6a7a8a;margin-bottom:1.2rem;">
            Paste any section of your manuscript and receive a rewritten version that
            maintains your scholarly argument while replacing generic, formulaic, or
            AI-generated language with an authentic academic voice.
        </div>
        """,
        unsafe_allow_html=True,
    )

    st.info(
        "**What this does:** This tool rewrites academic text to sound like the work of a "
        "confident, experienced HPE researcher — specific, evidence-grounded, and free of "
        "hollow phrases. It is benchmarked against published papers in JDE, EJDE, Medical Teacher, "
        "and Medical Education. Your argument and findings are preserved exactly."
    )

    col_ex1, col_ex2 = st.columns(2)
    with col_ex1:
        st.markdown(
            "<div style='background:#fdf0f0;border-left:3px solid #7a2020;"
            "padding:0.7rem 0.9rem;border-radius:0 3px 3px 0;font-size:0.82rem;"
            "font-family:Source Serif 4,serif;color:#4a1010;'>"
            "<strong>Before (generic):</strong><br>"
            "\"This study contributes to the existing literature by providing valuable "
            "insights into the use of AI in dental education. It is worth noting that "
            "the findings have important implications for future research and practice.\""
            "</div>",
            unsafe_allow_html=True,
        )
    with col_ex2:
        st.markdown(
            "<div style='background:#f0f8f4;border-left:3px solid #1d6b52;"
            "padding:0.7rem 0.9rem;border-radius:0 3px 3px 0;font-size:0.82rem;"
            "font-family:Source Serif 4,serif;color:#0f3a28;'>"
            "<strong>After (humanised):</strong><br>"
            "\"These findings extend our understanding of AI adoption in dental curricula "
            "by demonstrating that student familiarity with system descriptions — not "
            "system use itself — drives learning outcomes, a distinction prior studies "
            "have not isolated.\""
            "</div>",
            unsafe_allow_html=True,
        )

    st.divider()

    hum_section = st.selectbox(
        "Which section are you working on?",
        ["Introduction", "Methods", "Results", "Discussion", "Abstract",
         "Conclusion", "Literature Review", "Other"],
    )

    hum_target = st.selectbox(
        "Target journal register",
        ["Journal of Dental Education (JDE)",
         "European Journal of Dental Education (EJDE)",
         "Medical Teacher",
         "Medical Education",
         "BMC Medical Education",
         "General HPE journal"],
    )

    hum_instruction = st.text_area(
        "Any specific instruction? (optional)",
        placeholder="e.g. Keep the first sentence as is. Make the limitation paragraph more specific. "
                    "Do not change any statistics or findings.",
        height=70,
    )

    hum_input = st.text_area(
        "Paste the text you want to humanise",
        placeholder="Paste your paragraph or section here...",
        height=220,
    )

    col_btn1, col_btn2 = st.columns([2, 1])
    with col_btn1:
        run_humanise = st.button(
            "✍️ Humanise This Text →",
            disabled=not hum_input.strip(),
            use_container_width=True,
        )
    with col_btn2:
        hum_mode = st.selectbox(
            "Mode", ["Polish", "Full rewrite"], label_visibility="collapsed"
        )

    if run_humanise and hum_input.strip():
        mode_instruction = (
            "Lightly polish the text — improve clarity, remove generic phrases, "
            "strengthen the academic voice — but preserve the original structure and most wording."
            if hum_mode == "Polish"
            else
            "Fully rewrite the text with an authentic academic voice — restructure sentences, "
            "replace all generic phrases, and make the writing specific and evidence-grounded — "
            "while preserving every factual claim, finding, and argument."
        )

        hum_system = (
            f"You are an expert academic writing editor specialising in health professions education "
            f"and dental education research. You help researchers improve their writing so it reads "
            f"like high-quality published work in {hum_target}. "
            f"Your benchmark is papers published in JDE, EJDE, Medical Teacher, and Medical Education — "
            f"writing that is precise, specific, evidence-grounded, appropriately hedged, and free of "
            f"formulaic or AI-generated language. "
            f"You NEVER change findings, statistics, citations, or the scholarly argument. "
            f"You replace hollow phrases ('it is worth noting', 'this study contributes to the literature', "
            f"'valuable insights', 'important implications') with specific, concrete language. "
            f"You favour active voice where appropriate. You maintain a consistent scholarly register. "
            f"You write as a senior HPE researcher, not as a generic assistant."
        )

        hum_user = (
            f"Section: {hum_section}\n"
            f"Target journal: {hum_target}\n"
            f"Mode: {mode_instruction}\n"
            + (f"Specific instruction: {hum_instruction}\n" if hum_instruction.strip() else "")
            + f"\nTEXT TO HUMANISE:\n{hum_input}\n\n"
            f"Return ONLY:\n"
            f"1. The rewritten text (no preamble, no commentary)\n"
            f"2. Then a blank line\n"
            f"3. Then a brief bulleted list (max 5 items) of what you changed and why, "
            f"prefixed with '---CHANGES---'"
        )

        with st.spinner("Rewriting with an authentic scholarly voice..."):
            try:
                resp = client.messages.create(
                    model=FALLBACK_MODEL,
                    max_tokens=2048,
                    system=hum_system,
                    messages=[{"role": "user", "content": hum_user}],
                )
                hum_result = resp.content[0].text

                if "---CHANGES---" in hum_result:
                    parts       = hum_result.split("---CHANGES---", 1)
                    rewritten   = parts[0].strip()
                    changes     = parts[1].strip()
                else:
                    rewritten   = hum_result.strip()
                    changes     = ""

                st.markdown("#### Rewritten Text")
                st.markdown(
                    f"<div style='background:#f8f5ef;border:1px solid #e0dbd0;"
                    f"border-radius:4px;padding:1.2rem 1.4rem;"
                    f"font-family:Source Serif 4,Georgia,serif;"
                    f"font-size:0.9rem;line-height:1.85;color:#1a1612;'>"
                    f"{rewritten.replace(chr(10), '<br>')}"
                    f"</div>",
                    unsafe_allow_html=True,
                )

                col_copy, col_dl = st.columns([1, 1])
                with col_copy:
                    if st.button("📋 Copy to clipboard", key="copy_humanised"):
                        st.code(rewritten)
                with col_dl:
                    # Build Word document
                    from docx import Document as _Doc
                    from docx.shared import Pt as _Pt
                    from docx.enum.text import WD_ALIGN_PARAGRAPH as _WAP
                    from io import BytesIO as _BytesIO
                    _doc = _Doc()
                    _doc.styles["Normal"].font.name  = "Calibri"
                    _doc.styles["Normal"].font.size  = _Pt(11)
                    _h = _doc.add_heading("DentEdTech™ — Humanised Writing", 0)
                    _h.alignment = _WAP.CENTER
                    _doc.add_paragraph(f"Section: {hum_section}  ·  Target journal: {hum_target}")
                    _doc.add_paragraph("")
                    _doc.add_heading("Rewritten Text", level=2)
                    for _para in rewritten.split("\n"):
                        if _para.strip():
                            _doc.add_paragraph(_para.strip())
                    if changes:
                        _doc.add_paragraph("")
                        _doc.add_heading("Changes Made", level=2)
                        for _line in changes.split("\n"):
                            _line = _line.strip().lstrip("-•* ")
                            if _line:
                                _doc.add_paragraph(_line, style="List Bullet")
                    _fp = _doc.sections[0].footer.paragraphs[0]
                    _fp.text = "DentEdTech™  ·  akhalifah@taibahu.edu.sa  ·  Taibah University"
                    _buf = _BytesIO()
                    _doc.save(_buf)
                    st.download_button(
                        "⬇️ Download as Word (.docx)",
                        data=_buf.getvalue(),
                        file_name=f"DentEdTech_Humanised_{hum_section}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                    )

                if changes:
                    with st.expander("📝 What was changed and why", expanded=True):
                        for line in changes.split("\n"):
                            line = line.strip()
                            if line.startswith(("-", "•", "*")):
                                st.markdown(f"- {line.lstrip('-•* ')}")
                            elif line:
                                st.markdown(line)

            except Exception as e:
                st.error(f"Could not process: {e}")

    st.divider()
    st.caption(
        "The humanised text preserves all your findings, arguments, and citations. "
        "Always review the output carefully before submitting. "
        "DentEdTech™ · akhalifah@taibahu.edu.sa"
    )
